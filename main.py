import os
import json
import glob
import fitz
import comtypes.client
import time
import shutil
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageOps
from PyPDF2 import PdfMerger


SETTINGS = "settings.json"
rootdir = os.path.abspath(os.path.dirname(__file__))

#Функция преобразует docx в pdf.
def word2pdf(docx_path, pdf_path):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = True
    time.sleep(3)
    doc=word.Documents.Open(docx_path) 
    doc.SaveAs(pdf_path, FileFormat=17) 
    doc.Close() 
    word.Visible = False

#Функция для предоставления выбора файла.
def choose_file(files, greeting) -> int:
    while(True):
        print(greeting)
        for idx, file in enumerate(map(lambda filename: filename.split("\\")[-1], files)):
            print(f"\t{idx+1}.{file}")
        try:
            output_file_idx = int(input("Введите номер: "))
        except ValueError:
            print("Похоже вы ввели не число, попробуйте ещё раз...")
            continue
        if(output_file_idx < 1 or output_file_idx > len(files)):
            print("Похоже вы ввели неправильное число, попробуйте ещё раз")
            continue
        else:
            return output_file_idx - 1

#Функция для извлечения изображений с pdf файла.
def read_pdf(filename, output_dir):
    zoom_x = 4.0  
    zoom_y = 4.0  
    mat = fitz.Matrix(zoom_x, zoom_y) 
    with fitz.open(filename) as pdf_file:
        for idx, page in enumerate(pdf_file):
            pix = page.get_pixmap(matrix=mat)
            name = f"barcode{idx}.png"
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples, "raw", "RGB", 0, -1)
            img = img.rotate(90, expand=1)
            img = ImageOps.mirror(img)
            img.save(os.path.join(output_dir, name))
        

#Функция для чтения настроек.
def read_json(filename):
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileExistsError:
        print("ОШИБКА. Файл настроек не найден")
        return None
    
#Функция для чтения документа.
def read_document(filename):
    document = Document(filename)
    table = document.tables[0] 
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return document, table


#Функция для создания временных папок.
def check_dirs(settings):
    should_exist = ["image_sequence_dir", "temp_docx", "temp_pdf"]
    shouldnt_exist = ["output_docx"]
    for dir in should_exist:
        if not os.path.exists(settings[dir]):
            os.mkdir(settings[dir])
    for dir in shouldnt_exist:
        if os.path.exists(settings[dir]):
            shutil.rmtree(settings[dir])


#Функция для удаления временных файлов.
def delete_temp(settings):
    shouldnt_exist = ["image_sequence_dir", "temp_docx", "temp_pdf"]
    for dir in shouldnt_exist:
        if os.path.exists(settings[dir]):
            shutil.rmtree(settings[dir])
    


def main():
    #Получаем абсолютный путь до скрипта и получаем все файлы в этой директории.
    root = os.path.abspath(os.getcwd()) + '\\'
    files_in_directory = filter(os.path.isfile, glob.glob(root + '*'))
    files_in_directory = sorted(files_in_directory, key=os.path.getmtime)
    files_in_directory = list(filter(lambda file: ".pdf" in file and "output" not in file, files_in_directory))
    #Предоставляем выбор для исходного pdf файла.
    idx = choose_file(files_in_directory, "Выберите исходный файл:")


    print("Читаю настройки.")
    #Получаем настройки из json.
    settings = read_json(os.path.join(rootdir, SETTINGS))
    #Удаляем временные файлы и директории и создаём пустые.
    delete_temp(settings)
    check_dirs(settings)
    
    output_path = os.path.join(rootdir, settings["image_sequence_dir"])
    print("Читаю pdf.")
    #Экспортируем изображения из pdf.
    read_pdf(files_in_directory[idx], output_path)


    image_files = os.listdir(os.path.join(rootdir, settings["image_sequence_dir"]))
    if not len(image_files):
        print("Картинки в pdf не найдены !")
        exit(0)


    document, table = read_document(settings["input_docx"]) 


    #Переменная указывает ровно ли штрихкоды заполнят лист A4.
    overflow = False
    if len(image_files) % 8 != 0:
        overflow = True

    #Указатель на строку, колонку и страницу.
    row_ptr = 0
    col_ptr = 0
    table_ptr = 0
    #Список абсолютных путей временных .docx файлов.
    documents = []
    for i in range(len(image_files)):
        if col_ptr == 2:
            row_ptr += 1
            col_ptr = 0
        if row_ptr == 4:
            row_ptr = 0
            col_ptr = 0
            document_name = f"{settings['temp_docx']}\\N{table_ptr}_{settings['output_docx']}"
            document_path = os.path.join(rootdir, document_name)
            documents.append(document_path)
            document.save(document_name)
            document, table = read_document(settings["input_docx"])
            table_ptr += 1
        #В каждой ячейке таблицы создаётся параграф и помещается изображение размером A7.
        image_path = os.path.join(settings["image_sequence_dir"], image_files[i])
        paragraph = table.cell(row_ptr, col_ptr).add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(settings["barcode_width"]))
        col_ptr += 1


    #Все дальнейшие действия - настоящие извращения, т.к. проблема находится на стороне word.
    #Даже Аллах не ведает почему таблицы переполняются при сохранении многостраничного .docx документа.
    #Поэтому происходит временное сохранение одностраничных .docx документов, конвертация в .pdf и соединение его в один файл.


    #Если изображения не до конца заполняют таблицы, то последнее изображение сохраняется.
    if overflow:
        document_name = f"{settings['temp_docx']}\\N{table_ptr}_{settings['output_docx']}"
        document_path = os.path.join(rootdir, document_name)
        documents.append(document_path)
        document.save(document_path)
 
    #Переконвертация временных .docx в .pdf.
    print("Сохраняю pdf.")
    pdf_files = []
    for document in documents:
        output_path = '\\'.join(document.split('\\')[:-2])
        output_name = document.split('\\')[-1].split('.')[0] + ".pdf"
        output_name = os.path.join(settings["temp_pdf"], output_name)
        output_path = os.path.join(output_path, output_name)
        pdf_files.append(output_path)
        word2pdf(document, output_path)
    
    #Соединение всех pdf в один.
    print("Слепляю pdf.")
    merger = PdfMerger()
    for file in pdf_files:
        merger.append(open(file, 'rb'))


    time = datetime.datetime.now()
    time = time.strftime("%Y.%m.%d_%H.%M.%S")
    name = settings["output_pdf"].split('.')[0] + time + ".pdf"
    with open(name, 'wb') as f:
        merger.write(f)
        merger.close()


    delete_temp(settings)

if __name__ == "__main__":
    print("Made in barsukland!\n\n")
    try:
        main()
    except KeyboardInterrupt:
        pass
